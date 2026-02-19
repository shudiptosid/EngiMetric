import io
from datetime import datetime
from typing import Dict, Any
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT


class ExportService:
    """Generate PDF, DOCX, and XLSX exports for quotations and proposals."""

    PRIMARY_COLOR = HexColor("#1e40af")
    SECONDARY_COLOR = HexColor("#374151")
    ACCENT_COLOR = HexColor("#3b82f6")
    LIGHT_BG = HexColor("#f3f4f6")

    def generate_pdf(self, quotation_data: Dict[str, Any]) -> bytes:
        """Generate a professional PDF quotation/proposal."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=25 * mm,
            leftMargin=25 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        # Custom styles
        styles.add(ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=24,
            textColor=self.PRIMARY_COLOR,
            spaceAfter=6,
        ))
        styles.add(ParagraphStyle(
            "SectionHeader",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=self.PRIMARY_COLOR,
            spaceBefore=16,
            spaceAfter=8,
            borderPadding=(0, 0, 4, 0),
        ))
        styles.add(ParagraphStyle(
            "BodyCustom",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            textColor=self.SECONDARY_COLOR,
        ))
        styles.add(ParagraphStyle(
            "SmallGrey",
            parent=styles["Normal"],
            fontSize=8,
            textColor=HexColor("#9ca3af"),
        ))
        styles.add(ParagraphStyle(
            "TotalStyle",
            parent=styles["Normal"],
            fontSize=16,
            textColor=self.PRIMARY_COLOR,
            alignment=TA_RIGHT,
            fontName="Helvetica-Bold",
        ))

        content = quotation_data.get("content", {})
        elements = []

        # Header
        elements.append(Paragraph("EngineerCost Pro", styles["CustomTitle"]))
        elements.append(Paragraph("Professional Project Quotation", styles["SmallGrey"]))
        elements.append(Spacer(1, 8))
        elements.append(HRFlowable(width="100%", thickness=2, color=self.ACCENT_COLOR))
        elements.append(Spacer(1, 16))

        # Project Info
        title = quotation_data.get("title", "Project Quotation")
        elements.append(Paragraph(title, styles["SectionHeader"]))

        info_data = [
            ["Client:", content.get("client_name", "N/A")],
            ["Date:", datetime.now().strftime("%B %d, %Y")],
            ["Valid Until:", content.get("valid_until", "30 days from issue")],
            ["Status:", quotation_data.get("status", "Draft").upper()],
        ]
        info_table = Table(info_data, colWidths=[80, 400])
        info_table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("TEXTCOLOR", (0, 0), (0, -1), self.PRIMARY_COLOR),
            ("TEXTCOLOR", (1, 0), (1, -1), self.SECONDARY_COLOR),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(info_table)
        elements.append(Spacer(1, 16))

        # Project Overview
        if content.get("project_overview"):
            elements.append(Paragraph("Project Overview", styles["SectionHeader"]))
            elements.append(Paragraph(content["project_overview"], styles["BodyCustom"]))
            elements.append(Spacer(1, 10))

        # Scope of Work
        if content.get("scope_of_work"):
            elements.append(Paragraph("Scope of Work", styles["SectionHeader"]))
            for item in content["scope_of_work"]:
                elements.append(Paragraph(f"• {item}", styles["BodyCustom"]))
            elements.append(Spacer(1, 10))

        # Cost Breakdown
        if content.get("cost_breakdown"):
            elements.append(Paragraph("Cost Breakdown", styles["SectionHeader"]))
            breakdown = content["cost_breakdown"]
            cost_data = [["Item", "Amount"]]
            for key, value in breakdown.items():
                if key != "total":
                    label = key.replace("_", " ").title()
                    cost_data.append([label, f"${value:,.2f}"])

            cost_table = Table(cost_data, colWidths=[350, 130])
            cost_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), self.PRIMARY_COLOR),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#ffffff"), self.LIGHT_BG]),
                ("LINEBELOW", (0, 0), (-1, 0), 1, self.PRIMARY_COLOR),
                ("LINEBELOW", (0, -1), (-1, -1), 1, self.ACCENT_COLOR),
            ]))
            elements.append(cost_table)
            elements.append(Spacer(1, 12))

        # Total
        total = quotation_data.get("total_cost", 0)
        elements.append(Paragraph(f"Total: ${total:,.2f}", styles["TotalStyle"]))
        elements.append(Spacer(1, 16))

        # Timeline
        if content.get("timeline"):
            elements.append(Paragraph("Timeline", styles["SectionHeader"]))
            elements.append(Paragraph(content["timeline"], styles["BodyCustom"]))
            elements.append(Spacer(1, 10))

        # Payment Terms
        if content.get("payment_terms"):
            elements.append(Paragraph("Payment Terms", styles["SectionHeader"]))
            elements.append(Paragraph(content["payment_terms"], styles["BodyCustom"]))
            elements.append(Spacer(1, 10))

        # Legal
        elements.append(Spacer(1, 20))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#d1d5db")))
        elements.append(Spacer(1, 6))
        elements.append(Paragraph(
            "This quotation is valid for 30 days from the date of issue. "
            "Prices are subject to change based on scope modifications. "
            "Generated by EngineerCost Pro.",
            styles["SmallGrey"],
        ))

        doc.build(elements)
        return buffer.getvalue()

    def generate_docx(self, quotation_data: Dict[str, Any]) -> bytes:
        """Generate a DOCX quotation document."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        content = quotation_data.get("content", {})

        # Title
        title = doc.add_heading("EngineerCost Pro", level=0)
        for run in title.runs:
            run.font.color.rgb = RGBColor(30, 64, 175)

        doc.add_paragraph("Professional Project Quotation").italic = True
        doc.add_paragraph("─" * 60)

        # Project Info
        title_text = quotation_data.get("title", "Project Quotation")
        doc.add_heading(title_text, level=1)

        info = doc.add_paragraph()
        info.add_run("Client: ").bold = True
        info.add_run(content.get("client_name", "N/A") + "\n")
        info.add_run("Date: ").bold = True
        info.add_run(datetime.now().strftime("%B %d, %Y") + "\n")
        info.add_run("Status: ").bold = True
        info.add_run(quotation_data.get("status", "Draft").upper())

        # Overview
        if content.get("project_overview"):
            doc.add_heading("Project Overview", level=2)
            doc.add_paragraph(content["project_overview"])

        # Scope
        if content.get("scope_of_work"):
            doc.add_heading("Scope of Work", level=2)
            for item in content["scope_of_work"]:
                doc.add_paragraph(item, style="List Bullet")

        # Cost Breakdown
        if content.get("cost_breakdown"):
            doc.add_heading("Cost Breakdown", level=2)
            breakdown = content["cost_breakdown"]
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Shading Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Item"
            hdr[1].text = "Amount"
            for key, value in breakdown.items():
                if key != "total":
                    row = table.add_row().cells
                    row[0].text = key.replace("_", " ").title()
                    row[1].text = f"${value:,.2f}"

        # Total
        total = quotation_data.get("total_cost", 0)
        total_para = doc.add_paragraph()
        total_run = total_para.add_run(f"\nTotal: ${total:,.2f}")
        total_run.bold = True
        total_run.font.size = Pt(16)
        total_run.font.color.rgb = RGBColor(30, 64, 175)

        # Timeline & Payment
        if content.get("timeline"):
            doc.add_heading("Timeline", level=2)
            doc.add_paragraph(content["timeline"])

        if content.get("payment_terms"):
            doc.add_heading("Payment Terms", level=2)
            doc.add_paragraph(content["payment_terms"])

        # Footer
        doc.add_paragraph("─" * 60)
        footer = doc.add_paragraph(
            "This quotation is valid for 30 days. Generated by EngineerCost Pro."
        )
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(156, 163, 175)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_xlsx(self, quotation_data: Dict[str, Any]) -> bytes:
        """Generate an XLSX cost breakdown spreadsheet."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        wb = Workbook()
        ws = wb.active
        ws.title = "Cost Breakdown"
        content = quotation_data.get("content", {})

        # Styles
        header_font = Font(name="Calibri", size=12, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1E40AF", end_color="1E40AF", fill_type="solid")
        total_font = Font(name="Calibri", size=14, bold=True, color="1E40AF")
        label_font = Font(name="Calibri", size=11, bold=True)
        border = Border(
            bottom=Side(style="thin", color="D1D5DB"),
        )

        # Title
        ws.merge_cells("A1:E1")
        ws["A1"] = "EngineerCost Pro – Project Quotation"
        ws["A1"].font = Font(name="Calibri", size=18, bold=True, color="1E40AF")

        ws["A2"] = quotation_data.get("title", "Quotation")
        ws["A2"].font = Font(name="Calibri", size=12, color="374151")

        ws["A3"] = f"Client: {content.get('client_name', 'N/A')}"
        ws["A4"] = f"Date: {datetime.now().strftime('%B %d, %Y')}"

        # Cost Breakdown Header
        row = 6
        headers = ["Item", "Description", "Quantity", "Rate", "Amount"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        # Cost rows
        row = 7
        if content.get("cost_breakdown"):
            for key, value in content["cost_breakdown"].items():
                if key != "total":
                    ws.cell(row=row, column=1, value=key.replace("_", " ").title())
                    ws.cell(row=row, column=5, value=float(value))
                    ws.cell(row=row, column=5).number_format = '$#,##0.00'
                    for col in range(1, 6):
                        ws.cell(row=row, column=col).border = border
                    row += 1

        # Total row
        row += 1
        ws.cell(row=row, column=4, value="TOTAL:").font = total_font
        ws.cell(row=row, column=5, value=float(quotation_data.get("total_cost", 0)))
        ws.cell(row=row, column=5).font = total_font
        ws.cell(row=row, column=5).number_format = '$#,##0.00'

        # Column widths
        ws.column_dimensions["A"].width = 25
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 12
        ws.column_dimensions["D"].width = 12
        ws.column_dimensions["E"].width = 18

        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()


export_service = ExportService()
